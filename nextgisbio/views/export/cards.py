# encoding: utf-8

import os
from datetime import datetime
from tempfile import NamedTemporaryFile
import csv
from nextgisbio.utils import csv_utf
import pypandoc
import pyramid.httpexceptions as exc
from pyramid import security
from pyramid.renderers import render_to_response
from pyramid.response import FileResponse
from pyramid.view import view_config
from nextgisbio.models import DBSession, Cards, Taxon
from nextgisbio.views.cards.views_jtable import cards_jtable_browse
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from nextgisbio.utils.DictionaryMissingValues import DictionaryMissingValues


@view_config(route_name='export_cards_table')
def cards_table(request):
    if not security.authenticated_userid(request):
        raise exc.HTTPForbidden()

    output_format = request.GET['format']

    result = cards_jtable_browse(request)

    if output_format not in _get_formats():
        raise exc.HTTPException('Format "' + output_format + '" not supported')

    file_object = NamedTemporaryFile(suffix='.' + output_format)
    content_type = ''
    content_disposition = ''

    if output_format == 'pdf':
        _make_pdf(result, file_object, request)
        content_type = 'application/pdf'
        content_disposition = 'attachment; filename="{}"'.format('cards.pdf')

    if output_format == 'csv':
        _make_csv(result, file_object)
        content_type = 'text/csv'
        content_disposition = 'attachment; filename="{}"'.format('cards.csv')

    if output_format == 'docx':
        _make_docx(result, file_object)
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        content_disposition = 'attachment; filename="{}"'.format('cards.docx')

    file_response = FileResponse(os.path.abspath(file_object.name), content_type=content_type)
    file_response.content_disposition = content_disposition
    return file_response


def _get_formats():
    return {
        'pdf': True,
        'docx': True,
        'rtf': True,
        'odt': True,
        'csv': True
    }


def _make_pdf(result, file_object, request):
    params = {
        'date': datetime.now().strftime('%Y-%m-%d'),
        'result': result
    }
    html = render_to_response('export/cards-export.mako', params, request=request)
    extra_args = (
        '--smart',
        '--standalone',
        '--latex-engine=xelatex',
        '-V', 'mainfont:Linux Libertine O',
        '-V', 'geometry:top=2cm, bottom=2cm, left=2cm, right=2cm'
    )
    pypandoc.convert(html.body, 'pdf', format='markdown', outputfile=file_object.name, extra_args=extra_args)


def _make_csv(result, file_object):
    if len(result['Records']) == 0:
        return True

    with open(file_object.name, 'wb') as csv_file:
        writer = csv_utf.UnicodeWriter(csv_file, dialect=csv.Dialect.delimiter, delimiter=';')

        keys_sorted = sorted(result['Records'][0])
        writer.writerow(keys_sorted)

        cards_rows = []
        for card in result['Records']:
            card_row = []
            for key in keys_sorted:
                if key in card:
                    card_value = card[key]
                else:
                    card_value = ''
                if isinstance(card_value, unicode):
                    card_value = card_value.replace(';', ' ')
                card_row.append(card_value)
            cards_rows.append(card_row)

        writer.writerows(cards_rows)


def _make_docx(result, file_object):
    document = Document()
    for card in result['Records']:
        p = document.add_paragraph(u'ИНФОРМАЦИОННАЯ КАРТОЧКА')
        p.alignment = 1

        p = document.add_paragraph(u'встреч растений и грибов, занесенных в Красную книгу')
        p.alignment = 1

        p = document.add_paragraph(u'Ханты-Мансийского автономного округа')
        p.alignment = 1

        p = document.add_paragraph(card['taxon__name'])
        p.alignment = 1

        p = document.add_paragraph(card['taxon__russian_name'])
        p.alignment = 1

        p = document.add_paragraph(u'№ ' + str(card['cards__id']))
        p.alignment = 1

        card_for_rendering = DictionaryMissingValues(card)
        _add_paragraph(document, u'Место находки: ', card_for_rendering.get_value('cards__location', u'не описано'))
        _add_paragraph(document, u'Долгота: ', str(card['cards__lon']) if card['cards__lon'] else u'информации нет')
        _add_paragraph(document, u'Широта: ', str(card['cards__lat']) if card['cards__lat'] else u'информации нет')
        _add_paragraph(document, u'Местообитание: ', card['cards__habitat'] if card['cards__habitat'] else u'информации нет')
        _add_paragraph(document, u'Антропогенная нагрузка: ', card['cards__anthr_press'] if card['cards__anthr_press'] else u'информации нет')
        _add_paragraph(document, u'Лимитирующие факторы: ', card['cards__limit_fact'] if card['cards__limit_fact'] else u'информации нет')
        _add_paragraph(document, u'Меры защиты: ', card['cards__protection'] if card['cards__protection'] else u'информации нет')
        _add_paragraph(document, u'Фаза жизненного цикла: ', str(card['cards__pheno']) if card['cards__pheno'] else u'информации нет')
        _add_paragraph(document, u'Количество: ', str(card['cards__quantity']) if card['cards__quantity'] else u'информации нет')
        _add_paragraph(document, u'Площадь популяции: ', card['cards__area'] if card['cards__area'] else u'информации нет')
        _add_paragraph(document, u'Примечания: ', card['cards__notes'] if card['cards__notes'] else u'примечаний нет')
        _add_paragraph(document, u'Музейные образцы: ', card['cards__museum'] if card['cards__museum'] else u'информации нет')
        _add_paragraph(document, u'Источник информации: ', card['cards__publications'] if card['cards__publications'] else u'информации нет')
        _add_paragraph(document, u'Наблюдал: ', card_for_rendering.get_value('observer__name', u'информации нет'))

        document.add_page_break()

    document.save(file_object.name)


def _add_paragraph(document, key=None, value=None, alignment=0):
    p = document.add_paragraph(u'')
    p.alignment = alignment

    if key:
        p.add_run(key).bold = True

    if value:
        p.add_run(value)
